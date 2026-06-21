from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "user-guides" / "nofaclean-admin-system-guide-en-ar.docx"

BLUE = RGBColor(30, 64, 175)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
LIGHT_BLUE = "DBEAFE"
LIGHT_GRAY = "F3F4F6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_run_font(run, *, rtl: bool = False, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = run.font.size or Pt(10.5)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Arial")
    fonts.set(qn("w:cs"), "Arial")
    if rtl:
        rtl_node = OxmlElement("w:rtl")
        rtl_node.set(qn("w:val"), "1")
        r_pr.append(rtl_node)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size, color in (
        ("Title", 24, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, DARK),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_title(document: Document, text: str, *, subtitle: str | None = None, rtl: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    run.font.size = Pt(24)
    set_run_font(run, rtl=rtl, bold=True, color=BLUE)
    if subtitle:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_rtl(paragraph)
        run = paragraph.add_run(subtitle)
        run.font.size = Pt(11)
        set_run_font(run, rtl=rtl, color=MUTED)


def add_heading(document: Document, text: str, level: int = 1, *, rtl: bool = False) -> None:
    paragraph = document.add_heading("", level=level)
    if rtl:
        set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, rtl=rtl, bold=True, color=BLUE if level <= 2 else DARK)


def add_para(document: Document, text: str, *, rtl: bool = False, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if rtl:
        set_paragraph_rtl(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, rtl=rtl, bold=True, color=DARK)
        run = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(run, rtl=rtl)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, rtl=rtl)


def add_bullets(document: Document, items: Iterable[str], *, rtl: bool = False) -> None:
    for item in items:
        paragraph = document.add_paragraph()
        if rtl:
            set_paragraph_rtl(paragraph)
        run = paragraph.add_run(("• " if not rtl else "• ") + item)
        set_run_font(run, rtl=rtl)


def add_numbered(document: Document, items: Iterable[str], *, rtl: bool = False) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        if rtl:
            set_paragraph_rtl(paragraph)
        run = paragraph.add_run(f"{index}. {item}")
        set_run_font(run, rtl=rtl)


def add_table(document: Document, headers: list[str], rows: list[list[str]], *, rtl: bool = False) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        if rtl:
            set_paragraph_rtl(paragraph)
        run = paragraph.add_run(header)
        set_run_font(run, rtl=rtl, bold=True, color=DARK)

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cell = cells[col_index]
            set_cell_margins(cell)
            if row_index % 2:
                set_cell_shading(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            if rtl:
                set_paragraph_rtl(paragraph)
            run = paragraph.add_run(value)
            set_run_font(run, rtl=rtl)

    document.add_paragraph()


def add_callout(document: Document, title: str, body: str, *, rtl: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EFF6FF")
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    paragraph = cell.paragraphs[0]
    if rtl:
        set_paragraph_rtl(paragraph)
    title_run = paragraph.add_run(title + "\n")
    set_run_font(title_run, rtl=rtl, bold=True, color=BLUE)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, rtl=rtl)
    document.add_paragraph()


EN_SETUP_ROWS = [
    ["1", "Users and roles", "Create staff login accounts before daily work starts.", "Admin > Users"],
    ["2", "Workers", "Add workers, costs, status, skills, and availability notes.", "Workers"],
    ["3", "Services", "Define service pricing, workload, checklist, SLA/KPI, and optional packages.", "Services"],
    ["4", "Customers and sites", "Create the customer profile and exact service locations.", "Customers"],
    ["5", "Contracts", "Create the agreement from the customer, site, service, package, pricing, and terms.", "Contracts"],
    ["6", "Contract assignments", "Inside the contract details, schedule workers for specific days and tasks.", "Contracts > Details > Workers"],
    ["7", "Visits", "Generate visits and monitor today/tomorrow work from the operations board.", "Operations"],
    ["8", "Evidence review", "Review worker GPS, photos, checklist, issues, overtime, and quality results.", "Operations > Visits"],
    ["9", "Finance", "Issue invoices, review payment proofs, record collections, cheques, and credit notes.", "Finance"],
    ["10", "Expenses and reports", "Record actual company expenses and review profit, revenue, and worker reports.", "Expenses / Reports"],
]

AR_SETUP_ROWS = [
    ["١", "المستخدمون والصلاحيات", "إنشاء حسابات دخول الموظفين قبل تشغيل العمل اليومي.", "الإدارة > المستخدمون"],
    ["٢", "العمال", "إضافة بيانات العمال والتكلفة والحالة والمهارات وملاحظات التوفر.", "العمال"],
    ["٣", "الخدمات", "تعريف سعر الخدمة والجهد والقائمة التشغيلية ومؤشرات SLA/KPI والباقات الاختيارية.", "الخدمات"],
    ["٤", "العملاء والمواقع", "إنشاء ملف العميل ومواقع تقديم الخدمة بدقة.", "العملاء"],
    ["٥", "العقود", "إنشاء الاتفاقية من العميل والموقع والخدمة والباقة والسعر والشروط.", "العقود"],
    ["٦", "تعيينات العقد", "داخل تفاصيل العقد يتم جدولة العمال لأيام ومهام محددة.", "العقود > التفاصيل > العمال"],
    ["٧", "الزيارات", "توليد الزيارات ومتابعة عمل اليوم والغد من لوحة العمليات.", "العمليات"],
    ["٨", "مراجعة الأدلة", "مراجعة GPS والصور والقائمة والمشاكل والوقت الإضافي والجودة.", "العمليات > الزيارات"],
    ["٩", "المالية", "إصدار الفواتير ومراجعة إثباتات الدفع وتسجيل التحصيل والشيكات والإشعارات الدائنة.", "المالية"],
    ["١٠", "المصروفات والتقارير", "تسجيل المصروفات الفعلية ومراجعة الربحية والإيرادات وأداء العمال.", "المصروفات / التقارير"],
]


EN_ROLES = [
    ["Owner / General Manager", "Full access to the complete back office.", "Use for the business owner or top administrator only."],
    ["Operations", "Services, workers, customers, contracts, operations, expenses.", "Use for staff who plan visits, follow assignments, and manage daily work."],
    ["Supervisor", "Dashboard, workers, operations, worker portal.", "Use for field supervisors who review evidence and quality."],
    ["Accountant", "Dashboard, finance, expenses.", "Use for invoices, collections, cheques, VAT, and expense control."],
    ["Sales", "Services, customers, contracts.", "Use for staff who prepare offers and agreements."],
    ["Worker", "Worker mobile portal only.", "Use for cleaners and field workers who start/finish visits."],
    ["Customer", "Customer portal only.", "Use for customers who book services, accept contracts, upload payment proof, and rate visits."],
]

AR_ROLES = [
    ["المالك / المدير العام", "صلاحية كاملة لكل النظام الإداري.", "تستخدم لصاحب العمل أو المسؤول الأعلى فقط."],
    ["العمليات", "الخدمات، العمال، العملاء، العقود، العمليات، المصروفات.", "تستخدم لفريق التخطيط والمتابعة اليومية."],
    ["المشرف", "لوحة التحكم، العمال، العمليات، بوابة العامل.", "تستخدم للمشرفين الميدانيين لمراجعة الأدلة والجودة."],
    ["المحاسب", "لوحة التحكم، المالية، المصروفات.", "تستخدم للفواتير والتحصيل والشيكات والضريبة والمصروفات."],
    ["المبيعات", "الخدمات، العملاء، العقود.", "تستخدم لمن يجهز العروض والاتفاقيات."],
    ["العامل", "بوابة العامل فقط.", "تستخدم لعمال النظافة والميدان لتسجيل بداية ونهاية الزيارة."],
    ["العميل", "بوابة العميل فقط.", "تستخدم للعميل لطلب الخدمة وقبول العقد ورفع إثبات الدفع وتقييم الزيارة."],
]


def add_english_manual(document: Document) -> None:
    add_title(
        document,
        "Nofa Clean Admin System Guide",
        subtitle="English operator manual for nontechnical admin, operations, finance, and supervision teams.",
    )
    add_para(document, f"Document date: {date.today().isoformat()}")
    add_para(document, "System URL: https://nofaclean.com after production launch; local testing uses http://cleaning-nofacast.test.")
    add_callout(
        document,
        "Main rule",
        "Use the system in the correct order: prepare users, workers, services, customers, contracts, assignments, visits, evidence review, finance, expenses, then reports. Most errors happen when a later step is started before the required master data exists.",
    )

    add_heading(document, "1. Admin Workflow Overview")
    add_para(
        document,
        "The admin system is designed for cleaning and facility services where the business sells a service, optionally sells a package under that service, creates a customer contract, assigns workers to the contract schedule, monitors visits, reviews proof of work, invoices the customer, and measures profitability.",
    )
    add_numbered(
        document,
        [
            "Create staff users and assign the correct role.",
            "Add worker profiles, status, cost rate, skills, documents, and availability.",
            "Create services and optional packages. Services are the base configuration; packages are fixed offers under a service.",
            "Create customers and their service sites with contact details and map location.",
            "Create a contract using the customer, site, service, and package or custom agreement.",
            "Open the contract details and assign one or more workers for the correct day, time, role, and tasks.",
            "Generate visits and use Operations as the control center for today, tomorrow, alerts, requests, and visit review.",
            "Workers use the mobile portal to start work, submit checklist, upload photos, report issues, and finish work.",
            "Supervisors or admins review GPS, timing, photos, checklist, quality, issues, and overtime before approving completion.",
            "Finance issues invoices, approves payment proofs, records collections and cheques, and handles credit notes.",
            "Expenses records actual company costs. Reports show revenue, cost, gross profit, net profit, and worker performance.",
        ],
    )

    add_heading(document, "2. Recommended Setup Order")
    add_table(document, ["Step", "Create / review", "Why it must come first", "Page"], EN_SETUP_ROWS)

    add_heading(document, "3. Login and Language")
    add_bullets(
        document,
        [
            "Open the login page and sign in with the account provided by the admin.",
            "The system redirects staff to the admin dashboard, workers to the worker portal, and customers to the customer portal.",
            "Use the language switcher to change between English and Arabic. Layout direction changes automatically for Arabic.",
            "If a user cannot see a page, check their role and permissions in Users and Access.",
        ],
    )

    add_heading(document, "4. Users and Access")
    add_para(document, "Use this page to create staff accounts and control what every person is allowed to see or manage.")
    add_table(document, ["Role", "Access", "Recommended use"], EN_ROLES)
    add_heading(document, "Steps", level=3)
    add_numbered(
        document,
        [
            "Open Users from the admin sidebar.",
            "Click Create or add user.",
            "Enter the user's name, email, phone if available, password, and role.",
            "Select the correct role. Do not give Owner access unless the person should control everything.",
            "Save the user, then ask the staff member to log in and verify they can see the correct pages.",
        ],
    )
    add_callout(document, "Important", "If a user is both staff and customer, create the correct role for the work they need. Customer portal access should not expose admin pages.")

    add_heading(document, "5. Dashboard")
    add_para(document, "The dashboard is the daily business summary. It is not where detailed work is entered; it shows status, warnings, and business performance.")
    add_bullets(
        document,
        [
            "Operations cards show today visits, upcoming work, completion state, and attention items.",
            "Revenue and finance cards show collected amounts, open balances, overdue invoices, and expected collection.",
            "Worker performance shows attendance, completion, and productivity indicators.",
            "Use the dashboard each morning to decide which page needs action: Operations, Finance, Contracts, or Workers.",
        ],
    )

    add_heading(document, "6. Workers")
    add_para(document, "Workers are the people who perform service visits. Add workers before assigning contract schedules.")
    add_heading(document, "Worker list", level=3)
    add_bullets(
        document,
        [
            "Use search and filters to find workers by name, status, skill, or availability.",
            "Open Details to view performance, visits, targets, contracts, and compliance.",
            "Open Edit to update profile information, cost rate, documents, and status.",
            "Open Status / availability to review who is available before planning work.",
        ],
    )
    add_heading(document, "Worker form fields", level=3)
    add_table(
        document,
        ["Field", "Meaning", "How to use it"],
        [
            ["Employee code", "Internal code for the worker.", "Use a stable code from HR or payroll."],
            ["Name and phone", "Worker identity and contact.", "Required for coordination and supervisor follow-up."],
            ["Hired on", "Employment start date.", "Use for HR records and experience tracking."],
            ["Nationality and language", "Worker background and preferred language.", "Helps assign the right supervisor and instructions."],
            ["Job role", "Cleaner, supervisor, technician, driver, or other role.", "Use the real field role so assignments are clear."],
            ["Status", "Available, unavailable, inactive, or similar.", "Only available workers should be planned for new visits."],
            ["Cost rate", "Internal labor cost per hour or service calculation.", "Used for profitability and direct cost reports."],
            ["Availability notes", "Special timing notes or restrictions.", "Example: Fridays off, morning shift only, or assigned area."],
            ["Skills and certificates", "Service skills, training, compliance documents.", "Use for specialized services or customer requirements."],
        ],
    )
    add_heading(document, "Worker details tabs", level=3)
    add_table(
        document,
        ["Tab", "Purpose"],
        [
            ["Overview", "Basic profile, status, and current worker summary."],
            ["Performance", "Completion, attendance, quality, and productivity indicators."],
            ["Contracts", "Contracts where the worker is assigned."],
            ["Visits", "Past and upcoming visits for the worker."],
            ["Targets", "Worker goals and target tracking."],
            ["Compliance", "Documents, training, certificates, and readiness."],
        ],
    )

    add_heading(document, "7. Services and Packages")
    add_para(
        document,
        "A service is the base business offering, such as home cleaning, office cleaning, facility maintenance, car cleaning, or deep cleaning. A package is an optional fixed offer inside that service. Example: under Home Cleaning, packages can be Standard Weekly, Premium Monthly, or One-Time Deep Cleaning.",
    )
    add_callout(
        document,
        "Service vs Package",
        "If a contract uses no package, it uses the service/base/custom pricing. If it uses a package, the package price becomes the main price and the service defaults are inherited unless the admin customizes the terms.",
    )
    add_heading(document, "Service form fields", level=3)
    add_table(
        document,
        ["Field", "Meaning", "Guidance"],
        [
            ["Title", "Service name shown to admin and customers.", "Use clear names such as Residential Cleaning or Office Cleaning."],
            ["Category", "Service group.", "Choose cleaning, maintenance, car cleaning, or another relevant category."],
            ["Description", "Business explanation of the service.", "Keep it simple and customer friendly."],
            ["Pricing type", "How the service is priced.", "Monthly, fixed, hourly, or custom depending on the business offer."],
            ["Base price", "Default price in SAR before package changes.", "Use the normal starting price for the service."],
            ["VAT rate / prices include VAT", "Tax calculation setting.", "Use the Saudi VAT rules agreed by finance."],
            ["Default workers", "Normal worker count for one visit.", "This is copied into contracts unless changed."],
            ["Default duration", "Expected visit duration.", "Used for workload and capacity planning."],
            ["Visits per week / hours per visit", "Default frequency and time.", "Important for weekly and monthly contracts."],
            ["Extra hour rate", "Charge for approved overtime.", "Used when worker execution exceeds the agreed time."],
            ["Overtime policy", "How extra time is handled.", "Usually detected from actual worker start/finish time and approved by admin."],
            ["Material policy", "Included, chargeable, customer supplied, or mixed.", "Clarify who pays for consumables and equipment."],
            ["Material cost estimate", "Estimated service cost.", "Used for costing and profitability, not as an actual expense record."],
            ["Checklist", "Tasks workers must complete.", "Each line becomes a work instruction/check item."],
            ["SLA/KPI template", "Quality and performance targets.", "Use lines such as Attendance | 95 | percent | 40 | at_least | attendance."],
            ["Required certificates", "Compliance requirements.", "Use for chemicals, safety, or specialized work."],
            ["Active", "Whether the service is available.", "Deactivate services that should not be sold now."],
        ],
    )
    add_heading(document, "Package fields", level=3)
    add_table(
        document,
        ["Package field", "Purpose"],
        [
            ["Package title", "Name of the fixed offer."],
            ["Package price", "Main price when selected in a contract."],
            ["Worker count, visits per week, hours per visit", "Package workload."],
            ["Expected labor minutes", "Calculated workload for planning."],
            ["Package checklist", "Overrides or extends the service checklist."],
            ["Package SLA/KPI", "Overrides or extends the service KPI template."],
        ],
    )
    add_numbered(
        document,
        [
            "Create the service basics first: title, category, description, active state.",
            "Enter pricing: base price, VAT setting, minimum billable time if used.",
            "Enter workload defaults: workers, visit duration, visits per week, hours per visit.",
            "Enter material and equipment assumptions.",
            "Enter checklist and SLA/KPI template.",
            "Add optional packages only when the company sells fixed offers under the service.",
            "Save, then test the service by creating a draft contract to confirm the inherited values look right.",
        ],
    )

    add_heading(document, "8. Customers and Sites")
    add_para(document, "A customer can have one or more sites. A contract must be connected to a customer and a specific site.")
    add_table(
        document,
        ["Field", "Meaning", "Guidance"],
        [
            ["Customer type", "Individual, company, or other customer category.", "Choose the real billing/customer type."],
            ["Name", "Customer or company name.", "Use the name that should appear on contracts and invoices."],
            ["Phone and email", "Main communication details.", "Required for confirmations, payment follow-up, and portal use."],
            ["Preferred channel", "Phone, WhatsApp, email, or other.", "Use the channel the customer prefers."],
            ["Preferred language", "Arabic or English.", "Affects customer communication and document language expectations."],
            ["VAT number", "Tax registration number if applicable.", "Required for business invoices when available."],
            ["Status", "Active, inactive, lead, or similar.", "Only active customers should receive live contracts."],
            ["Country, city, district", "Structured location.", "Use Saudi Arabia locations for Nofa Clean Saudi operations."],
            ["Address", "Written site address.", "Add building, street, floor, and landmarks if needed."],
            ["Map pin / latitude / longitude", "Exact service location.", "Use the map pin where possible for worker navigation and GPS evidence."],
            ["Contact name and phone", "Site contact person.", "Important when the payer and site contact are different."],
            ["Default site", "Main site for the customer.", "Set the most used site as default."],
        ],
    )
    add_heading(document, "Customer details tabs", level=3)
    add_table(
        document,
        ["Tab", "Purpose"],
        [
            ["Overview", "Main customer profile and account summary."],
            ["Sites", "All service locations for this customer."],
            ["Contracts", "Customer agreements and contract status."],
            ["Workers", "Workers assigned to customer contracts."],
            ["Finance", "Invoices, paid amounts, open balance, and payment history."],
        ],
    )

    add_heading(document, "9. Contracts")
    add_para(
        document,
        "A contract is the official agreement between Nofa Clean and the customer. It connects customer, site, service, package or custom pricing, visit frequency, payment plan, terms, SLA/KPI, scope, and addendums.",
    )
    add_heading(document, "Contract creation steps", level=3)
    add_table(
        document,
        ["Step", "What to enter", "Notes"],
        [
            ["Customer and service", "Select customer, site, service, and optional package.", "Package list is filtered by selected service."],
            ["Terms", "Start date, end date, monthly fee, VAT, billing cycle, notice days, renewal, special terms.", "Terms are inherited from service/package until Customize terms is used."],
            ["Pricing and scope", "Workers, visits per week, hours per visit, planned weekly minutes, material policy, overtime policy, service scope.", "Use the real agreed scope, not only a sales estimate."],
            ["Payment plan", "Installment label, day, and percent.", "Total percent must equal 100%."],
            ["Addendums", "Optional numbered additions with title, summary, and effective date.", "Use for special site rules or later changes."],
            ["Review", "Check totals, VAT, material estimate, extra charges, package source, and custom agreement flag.", "Save only after reviewing the complete breakdown."],
        ],
    )
    add_callout(
        document,
        "Custom agreement",
        "When inherited terms are changed, the contract should be clearly understood as a customized agreement. This helps staff know the contract does not exactly follow the original service or package defaults.",
    )
    add_heading(document, "Contract details tabs", level=3)
    add_table(
        document,
        ["Tab", "What it does"],
        [
            ["Overview", "Shows agreement summary, customer, site, service, package, dates, status, and commercial summary."],
            ["Workers", "The correct place to assign one or more workers to the contract schedule. Choose main/support role, day, time, tasks, and instructions."],
            ["SLA / KPI", "Shows weekly/monthly service-level reports generated from visits, checklist, attendance, issues, photos, and supervisor inspections."],
            ["Payments", "Shows invoices, payment history, open balance, and related finance information."],
            ["Acceptance", "Shows customer acceptance, signature, customer decisions, and change requests."],
            ["Scope", "Shows service scope, tasks, terms, addendums, and work details."],
        ],
    )
    add_heading(document, "Assigning workers inside contract details", level=3)
    add_numbered(
        document,
        [
            "Open the contract details page.",
            "Open the Workers tab.",
            "Choose the date or weekday for the planned work.",
            "Select one or more workers.",
            "Set one worker as Main worker. The main worker is responsible for instructions and mobile tracking when applicable.",
            "Set other workers as Support workers.",
            "Enter start time, end time, and tasks or shared tasks.",
            "Use previous worker only when the same team should repeat the same schedule.",
            "Save the assignment.",
            "Review Operations to confirm the upcoming visit now has the correct worker planning.",
        ],
    )

    add_heading(document, "10. Operations")
    add_para(
        document,
        "Operations is the control center for the admin and supervisor team. It is not the main place to create contract assignments. Assignments should be managed inside the contract details. Operations shows what is scheduled, what needs attention, what requests are waiting, and which completed visits need review.",
    )
    add_table(
        document,
        ["Tab", "Purpose", "Main actions"],
        [
            ["Command", "Today/tomorrow control queue for contracts, alerts, payment attention, and visits needing acknowledgement.", "Open contract, open payment view, review alerts, follow up urgent actions."],
            ["Requests", "Customer booking requests, reschedule requests, and support tickets.", "Approve booking, reject booking, approve reschedule, mark in review, resolve, reject."],
            ["Follow-up", "Contracts coming this week that still need assignment or supervisor attention.", "Open contract and add workers/tasks in the contract details."],
            ["Capacity", "Worker availability lanes and conflicts for a selected date.", "Check who is free before assigning workers."],
            ["Schedule", "Weekly calendar and recurring assignment map.", "Generate visits and review the planned schedule."],
            ["Visits", "Daily visit monitoring and completed-visit review.", "Review evidence, approve completion, request correction, record issue, mark missed, approve overtime."],
        ],
    )
    add_heading(document, "Daily operations process", level=3)
    add_numbered(
        document,
        [
            "Open Operations every morning.",
            "Check Command for today/tomorrow contracts, payment attention, missed work, late work, and visits waiting for acknowledgement.",
            "Open Requests and process customer booking, reschedule, and support requests.",
            "Open Follow-up to find contracts that need workers before the next visit.",
            "Open the contract details from Follow-up and assign workers in the Workers tab.",
            "Use Capacity to confirm workers are available and not double-booked.",
            "Use Schedule to generate visits when needed and confirm upcoming work exists.",
            "Use Visits during the day to monitor status and after completion to review evidence.",
            "Approve completion only after checking timing, GPS, photos, checklist, issues, and customer feedback.",
        ],
    )
    add_callout(
        document,
        "Check-in and check-out",
        "Normally, the worker performs start/finish actions in the worker mobile portal. Admin check-in/check-out tools are administrative backup actions and should be used only when operations needs to correct or support a visit.",
    )
    add_heading(document, "Evidence and quality review", level=3)
    add_bullets(
        document,
        [
            "Photos prove site condition and work completion.",
            "GPS check-in/check-out proves the worker was at the location.",
            "Checklist completion proves the required scope was followed.",
            "Issue notes explain exceptions such as locked site, customer change, missing material, or damage.",
            "Overtime is detected from actual worker execution time and becomes billable only after admin approval.",
            "Customer ratings and low-rating follow-up help protect quality and retention.",
        ],
    )

    add_heading(document, "11. Finance")
    add_para(document, "Finance handles invoices, receivables, payments, payment proofs, cheques, credit notes, and billable overtime.")
    add_table(
        document,
        ["Area", "What it means", "Action"],
        [
            ["Receivables aging", "Open balances grouped by due date age.", "Use to follow overdue customers."],
            ["Expected collections", "Invoices expected in the next period.", "Use for cash planning."],
            ["Billable extras", "Approved overtime or extras waiting to be invoiced.", "Create invoice line when finance is ready."],
            ["Payment proof review", "Customer-uploaded transfer or payment evidence.", "View proof, verify externally, approve or reject with note."],
            ["Invoices", "Customer invoice rows with totals, paid, credit, balance, status.", "Record payment, update status, print invoice, create credit note."],
            ["Cheques", "Customer cheques and their state.", "Record cheque, then clear, bounce, or return it."],
            ["Credit notes", "Invoice reductions.", "Create and approve only after finance authorization."],
        ],
    )
    add_heading(document, "Payment proof approval", level=3)
    add_numbered(
        document,
        [
            "Open Finance.",
            "Go to Payment proof review.",
            "Open View proof and compare amount, date, customer, invoice, and reference with the bank or receipt.",
            "Click Approve proof only after verification.",
            "Click Reject if the proof is wrong, duplicated, unreadable, or not received by the company.",
            "Add an admin note so the customer or staff understands the decision.",
        ],
    )

    add_heading(document, "12. Expenses")
    add_para(
        document,
        "Expenses are actual company costs such as salaries, rent, bills, purchases, materials, transportation, and equipment. They affect net profit reports. Service material/equipment cost estimates are different; they are planning assumptions under services.",
    )
    add_table(
        document,
        ["Page", "Purpose"],
        [
            ["Expense records", "Enter actual expense transactions with date, category, type, vendor, amount, payment method, and notes."],
            ["Expense categories/types", "Manage the structure used to classify expenses."],
            ["Cost items/equipment/materials", "Maintain the cost catalog for materials and equipment."],
            ["Service cost links", "Link cost items to services for estimated service costing and profitability planning."],
        ],
    )

    add_heading(document, "13. Reports")
    add_para(document, "Reports are used by management to understand performance, cost, and profit.")
    add_bullets(
        document,
        [
            "Revenue report: sales and invoice performance.",
            "Direct costs report: labor/material cost related to delivered service.",
            "Company expenses report: actual overhead and operating expense.",
            "Gross profit report: revenue minus direct service costs.",
            "Net profit report: gross profit minus company expenses.",
            "Worker performance report: attendance, completion, quality, and productivity.",
            "Use Export when management needs CSV data outside the system.",
        ],
    )

    add_heading(document, "14. Settings")
    add_para(document, "Settings controls company-level configuration. Only authorized staff should change this page.")
    add_bullets(
        document,
        [
            "Company name and contact details affect documents and system identity.",
            "Locale and language settings affect English/Arabic behavior.",
            "Finance-related settings affect invoice behavior and should be checked by finance before launch.",
            "Do not change production settings during daily operations unless the change is approved.",
        ],
    )

    add_heading(document, "15. Related Customer and Worker Portals")
    add_table(
        document,
        ["Portal", "Who uses it", "Main functions"],
        [
            ["Customer portal", "Registered customers.", "Book service, view booking status, accept/sign contract, view visits, request reschedule/support, view invoices, upload payment proof, rate completed visit, manage profile/sites, read notifications."],
            ["Worker mobile portal", "Workers and field staff.", "See today's visits, start work, send GPS evidence, complete checklist, upload photos, report issues, record material usage and overtime, finish work."],
        ],
    )

    add_heading(document, "16. Button and Status Glossary")
    add_table(
        document,
        ["Button / status", "Meaning"],
        [
            ["Create / New", "Add a new record."],
            ["Edit", "Change an existing record."],
            ["Save", "Store the entered information."],
            ["View / Details", "Open the full page for a record."],
            ["Print / Download", "Open printable/PDF contract or invoice output."],
            ["Approve", "Accept a request, proof, completion, credit note, or overtime."],
            ["Reject", "Decline a request or proof with a reason when needed."],
            ["In review", "The request is being checked and is not yet resolved."],
            ["Resolve", "Close a support request after action has been completed."],
            ["Generate visits", "Create planned visit records from contract schedule."],
            ["Mark missed", "Record that a scheduled visit did not happen."],
            ["Request correction", "Ask the worker/supervisor to fix evidence or quality issue before final approval."],
            ["Main worker", "The lead worker responsible for the visit."],
            ["Support worker", "Additional worker assigned to the same visit."],
        ],
    )

    add_heading(document, "17. Common Mistakes to Avoid")
    add_bullets(
        document,
        [
            "Do not create a contract before the customer, site, and service are correct.",
            "Do not leave a package selected if the agreement is custom and should not use package price.",
            "Do not customize inherited contract terms without reviewing the custom agreement flag and price breakdown.",
            "Do not let payment plan percentages total anything other than 100%.",
            "Do not assign workers from the Operations board as the normal workflow; open the contract details and use the Workers tab.",
            "Do not approve visit completion until photos, GPS, checklist, issues, and timing have been reviewed.",
            "Do not approve payment proof before finance confirms the payment is actually received.",
            "Do not record service material estimates as actual company expenses. Actual expenses belong in Expenses.",
            "Do not give admin access to customer accounts or customer access to staff members unless there is a clear business reason.",
        ],
    )

    add_heading(document, "18. End-to-End UAT Checklist")
    add_numbered(
        document,
        [
            "Create one admin user, one operations user, one accountant, one supervisor, one worker, and one customer.",
            "Create a service with checklist, SLA/KPI template, material policy, and extra hour rule.",
            "Create at least one package under the service.",
            "Create a customer and site with country, city, district, address, and map coordinates.",
            "Create a contract using the package.",
            "Review the contract breakdown, payment plan, scope, addendums, and inherited SLA/KPI.",
            "Open the contract details and assign multiple workers, including one main worker.",
            "Generate visits and confirm the next visit appears in Operations.",
            "Log in as worker and complete a visit with checklist, photos, GPS evidence, issue if needed, and overtime if needed.",
            "Review and approve completion in Operations.",
            "Approve overtime and create invoice if applicable.",
            "Log in as customer, view contract/invoice, upload payment proof, and submit visit feedback.",
            "Approve payment proof in Finance and confirm invoice balance changes.",
            "Add a company expense.",
            "Check dashboard and reports for revenue, cost, expense, gross profit, and net profit.",
        ],
    )


def add_arabic_manual(document: Document) -> None:
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    add_title(
        document,
        "دليل النظام الإداري لنوفا كلين",
        subtitle="دليل عربي تشغيلي مبسط لفريق الإدارة والعمليات والمالية والإشراف.",
        rtl=True,
    )
    add_para(document, f"تاريخ المستند: {date.today().isoformat()}", rtl=True)
    add_para(document, "رابط النظام بعد الإطلاق: https://nofaclean.com، أما الاختبار المحلي فيكون على http://cleaning-nofacast.test.", rtl=True)
    add_callout(
        document,
        "القاعدة الأساسية",
        "استخدم النظام بالترتيب الصحيح: المستخدمون، العمال، الخدمات، العملاء، العقود، تعيينات العقد، الزيارات، مراجعة الأدلة، المالية، المصروفات، ثم التقارير. أغلب الأخطاء تحدث عند البدء بخطوة متقدمة قبل تجهيز البيانات الأساسية.",
        rtl=True,
    )

    add_heading(document, "١. فكرة سير العمل الإداري", rtl=True)
    add_para(
        document,
        "النظام الإداري مخصص لإدارة خدمات النظافة والمرافق. الفكرة تبدأ من تعريف الخدمة، ثم إنشاء باقة اختيارية، ثم إنشاء عقد للعميل، ثم تعيين العمال على جدول العقد، ثم متابعة الزيارات، ثم مراجعة إثبات العمل، ثم إصدار الفواتير وقياس الربحية.",
        rtl=True,
    )
    add_numbered(
        document,
        [
            "إنشاء حسابات الموظفين وتحديد الصلاحيات.",
            "إضافة ملفات العمال والحالة والتكلفة والمهارات والمستندات والتوفر.",
            "إنشاء الخدمات والباقات الاختيارية. الخدمة هي الأساس، والباقة عرض ثابت داخل الخدمة.",
            "إنشاء العملاء ومواقع تقديم الخدمة مع بيانات التواصل والموقع على الخريطة.",
            "إنشاء العقد من العميل والموقع والخدمة والباقة أو الاتفاق المخصص.",
            "فتح تفاصيل العقد وتعيين عامل أو أكثر حسب اليوم والوقت والدور والمهام.",
            "توليد الزيارات واستخدام العمليات كمركز متابعة لعمل اليوم والغد والتنبيهات والطلبات.",
            "العامل يستخدم بوابة الجوال لبدء العمل وتنفيذ القائمة ورفع الصور وتسجيل المشاكل وإنهاء العمل.",
            "المشرف أو المسؤول يراجع GPS والوقت والصور والقائمة والجودة والمشاكل والوقت الإضافي قبل اعتماد الإنجاز.",
            "المالية تصدر الفواتير وتراجع إثباتات الدفع وتسجل التحصيل والشيكات والإشعارات الدائنة.",
            "المصروفات تسجل التكاليف الفعلية للشركة، والتقارير تعرض الإيراد والتكلفة والربح وأداء العمال.",
        ],
        rtl=True,
    )

    add_heading(document, "٢. ترتيب الإعداد الموصى به", rtl=True)
    add_table(document, ["الخطوة", "ما يتم إنشاؤه / مراجعته", "لماذا يجب أن يأتي أولاً", "الصفحة"], AR_SETUP_ROWS, rtl=True)

    add_heading(document, "٣. تسجيل الدخول واللغة", rtl=True)
    add_bullets(
        document,
        [
            "افتح صفحة تسجيل الدخول وادخل بالحساب الذي يوفره المسؤول.",
            "النظام يوجه الموظف إلى لوحة الإدارة، والعامل إلى بوابة العامل، والعميل إلى بوابة العميل.",
            "استخدم زر اللغة للتبديل بين العربية والإنجليزية. اتجاه الصفحة يتغير تلقائياً عند اختيار العربية.",
            "إذا لم تظهر صفحة لمستخدم معين، راجع الدور والصلاحيات من صفحة المستخدمين والصلاحيات.",
        ],
        rtl=True,
    )

    add_heading(document, "٤. المستخدمون والصلاحيات", rtl=True)
    add_para(document, "تستخدم هذه الصفحة لإنشاء حسابات الموظفين والتحكم في الصفحات والعمليات التي يستطيع كل شخص الوصول إليها.", rtl=True)
    add_table(document, ["الدور", "الصلاحيات", "الاستخدام المناسب"], AR_ROLES, rtl=True)
    add_heading(document, "الخطوات", level=3, rtl=True)
    add_numbered(
        document,
        [
            "افتح صفحة المستخدمين من القائمة الجانبية.",
            "اضغط إنشاء أو إضافة مستخدم.",
            "أدخل الاسم والبريد والهاتف إن وجد وكلمة المرور والدور.",
            "اختر الدور الصحيح. لا تمنح صلاحية المالك إلا لمن يجب أن يتحكم في كل النظام.",
            "احفظ المستخدم، ثم اطلب منه تسجيل الدخول والتأكد من ظهور الصفحات المناسبة.",
        ],
        rtl=True,
    )
    add_callout(document, "تنبيه", "إذا كان الشخص عميلاً وموظفاً في نفس الوقت، استخدم الدور المناسب للعمل المطلوب. بوابة العميل لا يجب أن تعرض صفحات الإدارة.", rtl=True)

    add_heading(document, "٥. لوحة التحكم", rtl=True)
    add_para(document, "لوحة التحكم تعرض ملخص العمل اليومي. ليست لإدخال التفاصيل، بل لمعرفة الحالة والتنبيهات والأداء التجاري.", rtl=True)
    add_bullets(
        document,
        [
            "بطاقات العمليات تعرض زيارات اليوم والعمل القادم وحالة الإنجاز والعناصر التي تحتاج متابعة.",
            "بطاقات الإيراد والمالية تعرض المبالغ المحصلة والأرصدة المفتوحة والفواتير المتأخرة والتحصيل المتوقع.",
            "أداء العمال يعرض الحضور والإنجاز والجودة والإنتاجية.",
            "ابدأ يومك من لوحة التحكم لتعرف هل تحتاج إلى العمليات أو المالية أو العقود أو العمال.",
        ],
        rtl=True,
    )

    add_heading(document, "٦. العمال", rtl=True)
    add_para(document, "العمال هم الأشخاص الذين ينفذون الزيارات. يجب إضافتهم قبل جدولة العقود.", rtl=True)
    add_heading(document, "قائمة العمال", level=3, rtl=True)
    add_bullets(
        document,
        [
            "استخدم البحث والفلاتر للعثور على العامل حسب الاسم أو الحالة أو المهارة أو التوفر.",
            "افتح التفاصيل لرؤية الأداء والزيارات والأهداف والعقود والامتثال.",
            "افتح التعديل لتحديث الملف والتكلفة والمستندات والحالة.",
            "افتح حالة العمال / التوفر لمراجعة من هو متاح قبل التخطيط للعمل.",
        ],
        rtl=True,
    )
    add_heading(document, "حقول نموذج العامل", level=3, rtl=True)
    add_table(
        document,
        ["الحقل", "المعنى", "طريقة الاستخدام"],
        [
            ["رقم الموظف", "رمز داخلي للعامل.", "استخدم رمزاً ثابتاً من الموارد البشرية أو الرواتب."],
            ["الاسم والهاتف", "هوية العامل ووسيلة التواصل.", "ضروري للتنسيق والمتابعة."],
            ["تاريخ التعيين", "تاريخ بداية العمل.", "يستخدم لسجلات الموارد البشرية والخبرة."],
            ["الجنسية واللغة", "خلفية العامل واللغة المفضلة.", "يساعد في اختيار المشرف والتعليمات المناسبة."],
            ["الدور الوظيفي", "عامل نظافة، مشرف، فني، سائق أو غيره.", "استخدم الدور الحقيقي في الميدان."],
            ["الحالة", "متاح أو غير متاح أو غير نشط.", "لا تخطط زيارات جديدة إلا للعمال المتاحين."],
            ["معدل التكلفة", "تكلفة العامل داخلياً.", "تستخدم في الربحية والتكاليف المباشرة."],
            ["ملاحظات التوفر", "قيود أو ملاحظات وقتية.", "مثال: إجازة الجمعة، صباحاً فقط، أو منطقة محددة."],
            ["المهارات والشهادات", "مهارات الخدمة والتدريب والامتثال.", "تستخدم للخدمات المتخصصة أو متطلبات العميل."],
        ],
        rtl=True,
    )
    add_heading(document, "تبويبات تفاصيل العامل", level=3, rtl=True)
    add_table(
        document,
        ["التبويب", "الغرض"],
        [
            ["نظرة عامة", "البيانات الأساسية والحالة والملخص الحالي."],
            ["الأداء", "الإنجاز والحضور والجودة والإنتاجية."],
            ["العقود", "العقود التي تم تعيين العامل عليها."],
            ["الزيارات", "زيارات العامل السابقة والقادمة."],
            ["الأهداف", "أهداف العامل ومتابعة تحقيقها."],
            ["الامتثال", "المستندات والتدريب والشهادات والجاهزية."],
        ],
        rtl=True,
    )

    add_heading(document, "٧. الخدمات والباقات", rtl=True)
    add_para(
        document,
        "الخدمة هي العرض الأساسي للشركة مثل تنظيف المنازل أو تنظيف المكاتب أو صيانة المرافق أو تنظيف السيارات أو التنظيف العميق. الباقة هي عرض ثابت اختياري داخل الخدمة، مثل باقة أسبوعية أو شهرية.",
        rtl=True,
    )
    add_callout(
        document,
        "الخدمة مقابل الباقة",
        "إذا لم يستخدم العقد باقة، يعتمد على سعر الخدمة أو السعر المخصص. إذا استخدم باقة، يكون سعر الباقة هو السعر الأساسي، مع وراثة إعدادات الخدمة إلا إذا تم تخصيصها.",
        rtl=True,
    )
    add_heading(document, "حقول نموذج الخدمة", level=3, rtl=True)
    add_table(
        document,
        ["الحقل", "المعنى", "الإرشاد"],
        [
            ["اسم الخدمة", "الاسم الذي يظهر للإدارة والعملاء.", "استخدم اسماً واضحاً مثل تنظيف منازل أو تنظيف مكاتب."],
            ["التصنيف", "مجموعة الخدمة.", "اختر نظافة أو صيانة أو سيارات أو تصنيف مناسب."],
            ["الوصف", "شرح تجاري للخدمة.", "اجعله بسيطاً وواضحاً للعميل."],
            ["نوع التسعير", "طريقة تسعير الخدمة.", "شهري أو ثابت أو بالساعة أو مخصص حسب العرض."],
            ["السعر الأساسي", "السعر الافتراضي بالريال قبل تغييرات الباقات.", "استخدم سعر البداية المعتاد."],
            ["الضريبة / السعر شامل الضريبة", "إعداد احتساب الضريبة.", "استخدم إعدادات ضريبة القيمة المضافة المعتمدة في السعودية."],
            ["عدد العمال الافتراضي", "عدد العمال الطبيعي للزيارة.", "ينتقل إلى العقد إلا إذا تم تغييره."],
            ["مدة الزيارة الافتراضية", "الوقت المتوقع للزيارة.", "تستخدم في تخطيط القدرة والعمل."],
            ["عدد الزيارات أسبوعياً / ساعات الزيارة", "التكرار والوقت الافتراضي.", "مهم جداً للعقود الأسبوعية والشهرية."],
            ["سعر الساعة الإضافية", "رسوم الوقت الإضافي المعتمد.", "يستخدم عندما يتجاوز تنفيذ العامل الوقت المتفق عليه."],
            ["سياسة الوقت الإضافي", "طريقة التعامل مع الوقت الإضافي.", "غالباً يتم اكتشافه من وقت البداية والنهاية ثم يوافق عليه المسؤول."],
            ["سياسة المواد", "مشمولة أو مدفوعة أو من العميل أو مختلطة.", "وضح من يتحمل مواد التنظيف والمعدات."],
            ["تقدير تكلفة المواد", "تكلفة تقديرية للخدمة.", "تستخدم للتكلفة والربحية وليست مصروفاً فعلياً."],
            ["قائمة العمل", "مهام يجب على العامل تنفيذها.", "كل سطر يصبح مهمة أو عنصر تحقق."],
            ["قالب SLA/KPI", "أهداف مستوى الخدمة والأداء.", "مثال: الحضور | 95 | percent | 40 | at_least | attendance."],
            ["الشهادات المطلوبة", "متطلبات الامتثال.", "للخدمات التي تحتاج سلامة أو كيميائيات أو تخصص."],
            ["نشطة", "هل الخدمة متاحة للبيع.", "عطل الخدمات غير المتاحة حالياً."],
        ],
        rtl=True,
    )
    add_heading(document, "حقول الباقة", level=3, rtl=True)
    add_table(
        document,
        ["حقل الباقة", "الغرض"],
        [
            ["اسم الباقة", "اسم العرض الثابت."],
            ["سعر الباقة", "السعر الأساسي عند اختيارها في العقد."],
            ["عدد العمال والزيارات والساعات", "عبء العمل الخاص بالباقة."],
            ["دقائق العمل المتوقعة", "حساب التخطيط للجهد."],
            ["قائمة الباقة", "تعدل أو تضيف إلى قائمة الخدمة."],
            ["SLA/KPI للباقة", "تعدل أو تضيف إلى مؤشرات الخدمة."],
        ],
        rtl=True,
    )
    add_numbered(
        document,
        [
            "ابدأ ببيانات الخدمة الأساسية: الاسم والتصنيف والوصف والحالة.",
            "أدخل السعر والضريبة والحد الأدنى للفوترة إن وجد.",
            "أدخل الجهد الافتراضي: عدد العمال ومدة الزيارة وعدد الزيارات أسبوعياً وساعات الزيارة.",
            "أدخل افتراضات المواد والمعدات.",
            "أدخل قائمة العمل وقالب SLA/KPI.",
            "أضف الباقات فقط إذا كانت الشركة تبيع عروضاً ثابتة داخل الخدمة.",
            "احفظ، ثم اختبر الخدمة بإنشاء عقد مسودة للتأكد من انتقال القيم بشكل صحيح.",
        ],
        rtl=True,
    )

    add_heading(document, "٨. العملاء والمواقع", rtl=True)
    add_para(document, "يمكن للعميل أن يملك أكثر من موقع. يجب ربط كل عقد بعميل وموقع محدد.", rtl=True)
    add_table(
        document,
        ["الحقل", "المعنى", "الإرشاد"],
        [
            ["نوع العميل", "فرد أو شركة أو نوع آخر.", "اختر النوع الحقيقي للفوترة والعلاقة."],
            ["الاسم", "اسم العميل أو الشركة.", "استخدم الاسم الذي سيظهر في العقد والفاتورة."],
            ["الهاتف والبريد", "وسائل التواصل الأساسية.", "مهمة للتأكيدات والمتابعة وبوابة العميل."],
            ["قناة التواصل المفضلة", "هاتف أو واتساب أو بريد أو غيره.", "استخدم القناة التي يفضلها العميل."],
            ["اللغة المفضلة", "العربية أو الإنجليزية.", "تؤثر على التواصل وتوقعات المستندات."],
            ["الرقم الضريبي", "رقم تسجيل ضريبي إن وجد.", "مطلوب لفواتير الشركات عند توفره."],
            ["الحالة", "نشط أو غير نشط أو عميل محتمل.", "العقود الحية تكون للعملاء النشطين."],
            ["الدولة والمدينة والحي", "بيانات موقع منظمة.", "استخدم مواقع السعودية لعمليات نوفا كلين السعودية."],
            ["العنوان", "العنوان النصي للموقع.", "أضف المبنى والشارع والدور والمعالم."],
            ["دبوس الخريطة / خط العرض والطول", "الموقع الدقيق للخدمة.", "استخدم الخريطة لمساعدة العامل وإثبات GPS."],
            ["اسم وهاتف مسؤول الموقع", "الشخص الموجود في الموقع.", "مهم إذا كان الدافع مختلفاً عن مسؤول الموقع."],
            ["الموقع الافتراضي", "الموقع الرئيسي للعميل.", "اختر الموقع الأكثر استخداماً كافتراضي."],
        ],
        rtl=True,
    )
    add_heading(document, "تبويبات تفاصيل العميل", level=3, rtl=True)
    add_table(
        document,
        ["التبويب", "الغرض"],
        [
            ["نظرة عامة", "ملف العميل والملخص الرئيسي."],
            ["المواقع", "كل مواقع تقديم الخدمة للعميل."],
            ["العقود", "اتفاقيات العميل وحالاتها."],
            ["العمال", "العمال المرتبطون بعقود العميل."],
            ["المالية", "الفواتير والمدفوعات والرصيد المفتوح."],
        ],
        rtl=True,
    )

    add_heading(document, "٩. العقود", rtl=True)
    add_para(
        document,
        "العقد هو الاتفاق الرسمي بين نوفا كلين والعميل. يربط العميل والموقع والخدمة والباقة أو السعر المخصص وتكرار الزيارات وخطة الدفع والشروط وSLA/KPI والنطاق والملاحق.",
        rtl=True,
    )
    add_heading(document, "خطوات إنشاء العقد", level=3, rtl=True)
    add_table(
        document,
        ["الخطوة", "ما يتم إدخاله", "ملاحظات"],
        [
            ["العميل والخدمة", "اختيار العميل والموقع والخدمة والباقة الاختيارية.", "قائمة الباقات تتغير حسب الخدمة المختارة."],
            ["الشروط", "تاريخ البداية والنهاية والسعر الشهري والضريبة ودورة الفوترة وأيام الإشعار والتجديد والشروط الخاصة.", "تنتقل من الخدمة أو الباقة حتى يتم الضغط على تخصيص الشروط."],
            ["التسعير والنطاق", "عدد العمال والزيارات أسبوعياً وساعات الزيارة والدقائق المخططة وسياسة المواد والوقت الإضافي ونطاق الخدمة.", "استخدم النطاق المتفق عليه فعلياً وليس تقدير البيع فقط."],
            ["خطة الدفع", "اسم الدفعة واليوم والنسبة.", "يجب أن يكون مجموع النسب 100%."],
            ["الملاحق", "إضافات مرقمة بعنوان وملخص وتاريخ سريان.", "استخدمها لشروط الموقع الخاصة أو التغييرات اللاحقة."],
            ["المراجعة", "مراجعة الإجمالي والضريبة وتقدير المواد والرسوم الإضافية ومصدر الباقة وحالة التخصيص.", "لا تحفظ إلا بعد مراجعة التفصيل الكامل."],
        ],
        rtl=True,
    )
    add_callout(
        document,
        "اتفاق مخصص",
        "عند تعديل الشروط الموروثة، يجب فهم العقد كاتفاق مخصص. هذا يساعد الفريق على معرفة أن العقد لا يتبع إعدادات الخدمة أو الباقة كما هي.",
        rtl=True,
    )
    add_heading(document, "تبويبات تفاصيل العقد", level=3, rtl=True)
    add_table(
        document,
        ["التبويب", "ماذا يفعل"],
        [
            ["نظرة عامة", "يعرض ملخص الاتفاق والعميل والموقع والخدمة والباقة والتواريخ والحالة والملخص التجاري."],
            ["العمال", "المكان الصحيح لتعيين عامل أو أكثر على جدول العقد. اختر الدور الرئيسي أو المساند واليوم والوقت والمهام والتعليمات."],
            ["SLA / KPI", "يعرض تقارير أسبوعية وشهرية ناتجة من الزيارات والقائمة والحضور والمشاكل والصور وتفتيش المشرف."],
            ["المدفوعات", "يعرض الفواتير وتاريخ الدفع والرصيد والمعلومات المالية."],
            ["القبول", "يعرض قبول العميل والتوقيع وقرارات العميل وطلبات التغيير."],
            ["النطاق", "يعرض نطاق الخدمة والمهام والشروط والملاحق وتفاصيل العمل."],
        ],
        rtl=True,
    )
    add_heading(document, "تعيين العمال داخل تفاصيل العقد", level=3, rtl=True)
    add_numbered(
        document,
        [
            "افتح صفحة تفاصيل العقد.",
            "افتح تبويب العمال.",
            "اختر التاريخ أو يوم الأسبوع للعمل المخطط.",
            "اختر عامل واحد أو أكثر.",
            "حدد عاملاً واحداً كعامل رئيسي. العامل الرئيسي مسؤول عن التعليمات وتتبع الجوال عند الحاجة.",
            "حدد بقية العمال كعمال مساندين.",
            "أدخل وقت البداية ووقت النهاية والمهام أو المهام المشتركة.",
            "استخدم خيار العامل السابق فقط عندما يجب تكرار نفس الفريق.",
            "احفظ التعيين.",
            "راجع العمليات للتأكد من أن الزيارة القادمة أصبحت مخططة بالعمال الصحيحين.",
        ],
        rtl=True,
    )

    add_heading(document, "١٠. العمليات", rtl=True)
    add_para(
        document,
        "العمليات هي مركز التحكم لفريق الإدارة والإشراف. ليست المكان الرئيسي لإنشاء تعيينات العقود. التعيينات تتم من تفاصيل العقد، أما العمليات فتوضح ما هو مجدول وما يحتاج انتباه وما هي الطلبات المعلقة وما هي الزيارات المكتملة التي تحتاج مراجعة.",
        rtl=True,
    )
    add_table(
        document,
        ["التبويب", "الغرض", "الإجراءات الرئيسية"],
        [
            ["القيادة", "قائمة تحكم لليوم والغد تشمل العقود والتنبيهات ومتابعة الدفع والزيارات التي تحتاج اعتماد.", "فتح العقد، فتح المالية، متابعة التنبيهات والإجراءات العاجلة."],
            ["الطلبات", "طلبات الحجز من العملاء وطلبات إعادة الجدولة وتذاكر الدعم.", "قبول الحجز، رفض الحجز، قبول إعادة الجدولة، وضع تحت المراجعة، حل، رفض."],
            ["المتابعة", "عقود قادمة هذا الأسبوع تحتاج تعيين عمال أو متابعة مشرف.", "فتح العقد وإضافة العمال والمهام من تفاصيل العقد."],
            ["القدرة", "توفر العمال والتعارضات ليوم محدد.", "التأكد من أن العامل متاح قبل تعيينه."],
            ["الجدول", "تقويم أسبوعي وخريطة التعيينات المتكررة.", "توليد الزيارات ومراجعة الجدول المخطط."],
            ["الزيارات", "متابعة الزيارات اليومية ومراجعة الزيارات المكتملة.", "مراجعة الأدلة، اعتماد الإنجاز، طلب تصحيح، تسجيل مشكلة، تسجيل زيارة فائتة، اعتماد الوقت الإضافي."],
        ],
        rtl=True,
    )
    add_heading(document, "العملية اليومية لفريق العمليات", level=3, rtl=True)
    add_numbered(
        document,
        [
            "افتح العمليات كل صباح.",
            "راجع تبويب القيادة لمعرفة عقود اليوم والغد والتنبيهات والمدفوعات والزيارات التي تحتاج اعتماد.",
            "افتح الطلبات وعالج طلبات الحجز وإعادة الجدولة والدعم.",
            "افتح المتابعة لمعرفة العقود التي تحتاج عمالاً قبل الزيارة القادمة.",
            "افتح تفاصيل العقد من المتابعة وأضف العمال من تبويب العمال.",
            "استخدم القدرة للتأكد من توفر العمال وعدم وجود تعارض.",
            "استخدم الجدول لتوليد الزيارات عند الحاجة والتأكد من وجود عمل قادم.",
            "استخدم الزيارات أثناء اليوم للمتابعة وبعد الإنجاز لمراجعة الأدلة.",
            "لا تعتمد الإنجاز إلا بعد مراجعة الوقت وGPS والصور والقائمة والمشاكل وتقييم العميل.",
        ],
        rtl=True,
    )
    add_callout(
        document,
        "تسجيل البداية والنهاية",
        "في الوضع الطبيعي، العامل هو من يسجل بداية ونهاية العمل من بوابة العامل على الجوال. أدوات البداية والنهاية في الإدارة هي إجراءات مساندة أو تصحيحية وليست طريقة العمل اليومية الأساسية.",
        rtl=True,
    )
    add_heading(document, "مراجعة الأدلة والجودة", level=3, rtl=True)
    add_bullets(
        document,
        [
            "الصور تثبت حالة الموقع وإنجاز العمل.",
            "GPS عند البداية والنهاية يثبت أن العامل كان في الموقع.",
            "القائمة تثبت أن نطاق العمل المطلوب تم اتباعه.",
            "ملاحظات المشاكل تشرح الاستثناءات مثل الموقع مغلق أو تغيير من العميل أو نقص مواد أو ضرر.",
            "الوقت الإضافي يتم اكتشافه من وقت التنفيذ الفعلي ويصبح قابلاً للفوترة فقط بعد اعتماد الإدارة.",
            "تقييم العميل ومتابعة التقييم المنخفض تساعد على حماية الجودة والاحتفاظ بالعملاء.",
        ],
        rtl=True,
    )

    add_heading(document, "١١. المالية", rtl=True)
    add_para(document, "المالية تدير الفواتير والذمم والمدفوعات وإثباتات الدفع والشيكات والإشعارات الدائنة والوقت الإضافي القابل للفوترة.", rtl=True)
    add_table(
        document,
        ["القسم", "المعنى", "الإجراء"],
        [
            ["أعمار الذمم", "الأرصدة المفتوحة حسب عمر تاريخ الاستحقاق.", "تستخدم لمتابعة العملاء المتأخرين."],
            ["التحصيل المتوقع", "الفواتير المتوقع تحصيلها قريباً.", "تستخدم للتخطيط النقدي."],
            ["الإضافات القابلة للفوترة", "وقت إضافي أو إضافات معتمدة بانتظار الفوترة.", "إنشاء بند فاتورة عندما تكون المالية جاهزة."],
            ["مراجعة إثبات الدفع", "إثبات تحويل أو دفع مرفوع من العميل.", "فتح الإثبات والتحقق خارجياً ثم القبول أو الرفض مع ملاحظة."],
            ["الفواتير", "صفوف الفواتير مع الإجمالي والمدفوع والدائن والرصيد والحالة.", "تسجيل دفع، تحديث حالة، طباعة فاتورة، إنشاء إشعار دائن."],
            ["الشيكات", "شيكات العملاء وحالتها.", "تسجيل شيك ثم تحصيله أو إرجاعه أو تسجيله كمرتجع."],
            ["الإشعارات الدائنة", "تخفيضات على الفواتير.", "تنشأ وتعتمد فقط بعد موافقة مالية."],
        ],
        rtl=True,
    )
    add_heading(document, "اعتماد إثبات الدفع", level=3, rtl=True)
    add_numbered(
        document,
        [
            "افتح المالية.",
            "اذهب إلى مراجعة إثباتات الدفع.",
            "افتح الإثبات وقارن المبلغ والتاريخ والعميل والفاتورة والمرجع مع البنك أو الإيصال.",
            "اضغط اعتماد الإثبات فقط بعد التحقق.",
            "اضغط رفض إذا كان الإثبات خاطئاً أو مكرراً أو غير واضح أو لم تستلمه الشركة.",
            "أضف ملاحظة إدارية حتى يفهم العميل أو الموظف سبب القرار.",
        ],
        rtl=True,
    )

    add_heading(document, "١٢. المصروفات", rtl=True)
    add_para(
        document,
        "المصروفات هي تكاليف فعلية على الشركة مثل الرواتب والإيجار والفواتير والمشتريات والمواد والنقل والمعدات. تؤثر على تقارير صافي الربح. تقديرات مواد ومعدات الخدمة مختلفة؛ هي افتراضات تخطيط داخل الخدمة.",
        rtl=True,
    )
    add_table(
        document,
        ["الصفحة", "الغرض"],
        [
            ["سجلات المصروفات", "إدخال المصروفات الفعلية بالتاريخ والتصنيف والنوع والمورد والمبلغ وطريقة الدفع والملاحظات."],
            ["تصنيفات وأنواع المصروفات", "إدارة الهيكل المستخدم لتصنيف المصروفات."],
            ["عناصر التكلفة / المعدات / المواد", "إدارة كتالوج تكلفة المواد والمعدات."],
            ["ربط تكاليف الخدمة", "ربط عناصر التكلفة بالخدمات لتقدير تكلفة الخدمة والربحية."],
        ],
        rtl=True,
    )

    add_heading(document, "١٣. التقارير", rtl=True)
    add_para(document, "تستخدم التقارير للإدارة لفهم الأداء والتكلفة والربحية.", rtl=True)
    add_bullets(
        document,
        [
            "تقرير الإيرادات: المبيعات وأداء الفواتير.",
            "تقرير التكاليف المباشرة: تكلفة العمال والمواد المرتبطة بالخدمة المنفذة.",
            "تقرير مصروفات الشركة: المصروفات التشغيلية والفعلية.",
            "تقرير إجمالي الربح: الإيرادات ناقص التكاليف المباشرة.",
            "تقرير صافي الربح: إجمالي الربح ناقص مصروفات الشركة.",
            "تقرير أداء العمال: الحضور والإنجاز والجودة والإنتاجية.",
            "استخدم التصدير عندما تحتاج الإدارة إلى البيانات خارج النظام بصيغة CSV.",
        ],
        rtl=True,
    )

    add_heading(document, "١٤. الإعدادات", rtl=True)
    add_para(document, "الإعدادات تتحكم في تهيئة الشركة والنظام. لا يجب تعديلها إلا من موظف مخول.", rtl=True)
    add_bullets(
        document,
        [
            "اسم الشركة وبيانات التواصل تؤثر على المستندات وهوية النظام.",
            "إعدادات اللغة تؤثر على العربية والإنجليزية واتجاه الصفحة.",
            "الإعدادات المالية تؤثر على سلوك الفواتير ويجب مراجعتها من المالية قبل الإطلاق.",
            "لا تغير إعدادات الإنتاج أثناء التشغيل اليومي إلا بعد موافقة واضحة.",
        ],
        rtl=True,
    )

    add_heading(document, "١٥. بوابة العميل وبوابة العامل", rtl=True)
    add_table(
        document,
        ["البوابة", "من يستخدمها", "الوظائف الرئيسية"],
        [
            ["بوابة العميل", "العملاء المسجلون.", "طلب خدمة، متابعة حالة الطلب، قبول وتوقيع العقد، رؤية الزيارات، طلب إعادة جدولة أو دعم، رؤية الفواتير، رفع إثبات الدفع، تقييم الزيارة، إدارة الملف والمواقع، قراءة الإشعارات."],
            ["بوابة العامل", "العمال والفريق الميداني.", "رؤية زيارات اليوم، بدء العمل، إرسال إثبات GPS، تنفيذ القائمة، رفع الصور، تسجيل المشاكل، تسجيل المواد والوقت الإضافي، إنهاء العمل."],
        ],
        rtl=True,
    )

    add_heading(document, "١٦. قاموس الأزرار والحالات", rtl=True)
    add_table(
        document,
        ["الزر / الحالة", "المعنى"],
        [
            ["إنشاء / جديد", "إضافة سجل جديد."],
            ["تعديل", "تغيير سجل موجود."],
            ["حفظ", "تخزين المعلومات المدخلة."],
            ["عرض / تفاصيل", "فتح الصفحة الكاملة للسجل."],
            ["طباعة / تنزيل", "فتح نسخة قابلة للطباعة أو PDF للعقد أو الفاتورة."],
            ["اعتماد", "قبول طلب أو إثبات أو إنجاز أو إشعار دائن أو وقت إضافي."],
            ["رفض", "رفض طلب أو إثبات مع سبب عند الحاجة."],
            ["تحت المراجعة", "الطلب قيد الفحص ولم يتم حله بعد."],
            ["حل", "إغلاق طلب دعم بعد تنفيذ الإجراء."],
            ["توليد الزيارات", "إنشاء زيارات مخططة من جدول العقد."],
            ["تسجيل زيارة فائتة", "تسجيل أن الزيارة المجدولة لم تحدث."],
            ["طلب تصحيح", "طلب إصلاح دليل أو مشكلة جودة قبل الاعتماد النهائي."],
            ["عامل رئيسي", "العامل المسؤول عن الزيارة."],
            ["عامل مساند", "عامل إضافي معين على نفس الزيارة."],
        ],
        rtl=True,
    )

    add_heading(document, "١٧. أخطاء شائعة يجب تجنبها", rtl=True)
    add_bullets(
        document,
        [
            "لا تنشئ عقداً قبل التأكد من صحة العميل والموقع والخدمة.",
            "لا تترك باقة مختارة إذا كان الاتفاق مخصصاً ولا يجب أن يستخدم سعر الباقة.",
            "لا تخصص شروط العقد الموروثة دون مراجعة علامة الاتفاق المخصص وتفصيل السعر.",
            "لا تترك مجموع نسب خطة الدفع أقل أو أكثر من 100%.",
            "لا تجعل لوحة العمليات هي طريقة تعيين العمال اليومية؛ افتح تفاصيل العقد واستخدم تبويب العمال.",
            "لا تعتمد إنجاز الزيارة قبل مراجعة الصور وGPS والقائمة والمشاكل والوقت.",
            "لا تعتمد إثبات الدفع قبل تأكيد المالية أن المبلغ وصل فعلاً.",
            "لا تسجل تقديرات مواد الخدمة كمصروفات فعلية. المصروفات الفعلية تدخل في المصروفات.",
            "لا تمنح صلاحيات إدارة لحساب عميل أو صلاحيات عميل لموظف إلا لسبب تجاري واضح.",
        ],
        rtl=True,
    )

    add_heading(document, "١٨. قائمة اختبار شاملة قبل الاعتماد", rtl=True)
    add_numbered(
        document,
        [
            "أنشئ مستخدم إدارة ومستخدم عمليات ومحاسب ومشرف وعامل وعميل.",
            "أنشئ خدمة تحتوي على قائمة عمل وقالب SLA/KPI وسياسة مواد وقاعدة وقت إضافي.",
            "أنشئ باقة واحدة على الأقل داخل الخدمة.",
            "أنشئ عميلاً وموقعاً بالدولة والمدينة والحي والعنوان وإحداثيات الخريطة.",
            "أنشئ عقداً باستخدام الباقة.",
            "راجع تفصيل العقد وخطة الدفع والنطاق والملاحق وSLA/KPI الموروث.",
            "افتح تفاصيل العقد وعيّن عدة عمال، منهم عامل رئيسي واحد.",
            "ولّد الزيارات وتأكد من ظهور الزيارة القادمة في العمليات.",
            "سجل الدخول كعامل وأكمل زيارة بقائمة وصور وGPS ومشكلة إذا وجدت ووقت إضافي إذا وجد.",
            "راجع واعتمد الإنجاز من العمليات.",
            "اعتمد الوقت الإضافي وأنشئ فاتورة إذا لزم.",
            "سجل الدخول كعميل واعرض العقد والفاتورة وارفع إثبات الدفع وأرسل تقييم الزيارة.",
            "اعتمد إثبات الدفع من المالية وتأكد من تغير رصيد الفاتورة.",
            "أضف مصروفاً فعلياً للشركة.",
            "راجع لوحة التحكم والتقارير للإيراد والتكلفة والمصروفات وإجمالي الربح وصافي الربح.",
        ],
        rtl=True,
    )


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Nofa Clean")
    run.font.size = Pt(28)
    set_run_font(run, bold=True, color=BLUE)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Admin System Guide\nدليل النظام الإداري")
    run.font.size = Pt(22)
    set_run_font(run, bold=True, color=DARK)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("English + Arabic / إنجليزي + عربي")
    run.font.size = Pt(12)
    set_run_font(run, color=MUTED)

    document.add_paragraph()
    add_table(
        document,
        ["Item", "Details"],
        [
            ["Audience", "Admin, operations, finance, supervisors, and business review teams"],
            ["Purpose", "Step-by-step nontechnical guide for daily use of the Nofa Clean admin system"],
            ["System", "Laravel / Inertia / Vue cleaning and facility services platform"],
            ["Production URL", "https://nofaclean.com"],
            ["Version", f"UAT/admin guide draft - {date.today().isoformat()}"],
        ],
    )
    add_para(
        document,
        "This guide explains how to operate the admin system in the intended business order. It is written for nontechnical users and focuses on what to enter, where to enter it, and why each page matters.",
    )
    document.add_page_break()


def build() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)
    add_english_manual(document)
    add_arabic_manual(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
